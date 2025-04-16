from groq import Groq
from json import load, dump
import datetime
import subprocess
import os
from dotenv import dotenv_values
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

# Load environment variables
env_vars = dotenv_values(".env")
Username = env_vars.get("Username")
Assistantname = env_vars.get("Assistantname")
GroqAPIKey = env_vars.get("GroqAPIKey")

# Initialize Groq client
client = Groq(api_key=GroqAPIKey)

# Load chat history
try:
    with open(r"Data\ChatLog.json", "r") as f:
        messages = load(f)
except FileNotFoundError:
    messages = []

# System Prompt for AI
System = f"""Hello, I am {Username}, You are a very accurate and advanced AI chatbot named {Assistantname} which also has real-time up-to-date information from the internet.
*** Do not tell time until I ask, do not talk too much, just answer the question.***
*** Reply in only English, even if the question is in Hindi, reply in English.***
*** Do not provide notes in the output, just answer the question and never mention your training data. ***
"""

SystemChatBot = [{"role": "system", "content": System}]

# Function to get real-time information
def RealtimeInformation():
    now = datetime.datetime.now()
    return f"Please use this real-time information if needed,\nDay: {now.strftime('%A')}\nDate: {now.strftime('%d')}\nMonth: {now.strftime('%B')}\nYear: {now.strftime('%Y')}\nTime: {now.strftime('%H:%M:%S')}."

# Function to clean up AI responses
def AnswerModifier(Answer):
    return '\n'.join([line for line in Answer.split('\n') if line.strip()])

# Chatbot function to process user queries
def ChatBot(Query):
    global messages
    try:
        messages.append({"role": "user", "content": Query})

        completion = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=SystemChatBot + [{"role": "system", "content": RealtimeInformation()}] + messages,
            max_tokens=1024,
            temperature=0.7,
            top_p=1,
            stream=True
        )

        Answer = ""
        for chunk in completion:
            if chunk.choices[0].delta.content:
                Answer += chunk.choices[0].delta.content

        Answer = Answer.replace("</s>", "").strip()
        messages.append({"role": "assistant", "content": Answer})

        with open(r"Data\ChatLog.json", "w") as f:
            dump(messages, f, indent=4)

        return AnswerModifier(Answer)

    except Exception as e:
        print(f"Error: {e}")
        return "An error occurred while processing your request."

# Function to open Notepad (for text files)
def OpenNotepad(File):
    subprocess.Popen(["notepad.exe", File])

# AI Content Writer Function
def ContentWriterAI(prompt):
    messages.append({"role": "user", "content": prompt})

    completion = client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=SystemChatBot + messages,
        max_tokens=2048,
        temperature=0.7,
        top_p=1,
        stream=True
    )

    Answer = ""
    for chunk in completion:
        if chunk.choices[0].delta.content:
            Answer += chunk.choices[0].delta.content

    Answer = Answer.replace("</s>", "").strip()
    messages.append({"role": "assistant", "content": Answer})
    return Answer

# Function to create a professional resume with photo support
def CreateAdvancedResume(content, filename, photo_path=None):
    doc = Document()

    # Add Profile Picture
    if photo_path and os.path.exists(photo_path):
        doc.add_picture(photo_path, width=Inches(1.5))
    
    # Title: Resume
    title = doc.add_paragraph()
    title_run = title.add_run("Resume")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n")  # Line break

    # Formatting Sections
    sections = ["Summary", "Skills", "Experience", "Education", "Projects", "Certifications", "Contact"]
    content_lines = content.split("\n")
    section_idx = -1

    for line in content_lines:
        line = line.strip()
        if not line:
            continue

        # If line matches a section title, format it
        if any(line.lower().startswith(sec.lower()) for sec in sections):
            section_idx += 1
            para = doc.add_paragraph()
            run = para.add_run("\n" + line)
            run.bold = True
            run.font.size = Pt(12)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            para = doc.add_paragraph("• " + line)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_after = Pt(5)

    # Save Document
    filepath = rf"Data\{filename}.docx"
    doc.save(filepath)

    # Open in Microsoft Word
    os.startfile(filepath)

# Function to handle content writing (including resume generation)
def Content(Topic):
    Topic = Topic.strip().lower()

    if "resume" in Topic:
        ContentByAI = ContentWriterAI("Write a highly advanced professional resume with proper formatting.")
        filename = "Advanced_Resume"

        # Photo Path (Provide a valid photo path here)
        photo_path = "Data/profile_photo.jpg"  # Change this to the actual path of the photo

        CreateAdvancedResume(ContentByAI, filename, photo_path)
    else:
        ContentByAI = ContentWriterAI(Topic)
        filename = rf"Data\{Topic.replace(' ', '_')}.txt"
        with open(filename, "w", encoding="utf-8") as file:
            file.write(ContentByAI)
        OpenNotepad(filename)

    return True

if __name__ == "__main__":
    while True:
        user_input = input("Enter Your Question: ").strip().lower()

        content_keywords = [
            "write an application on sick leave",
            "write an application on half leave",
            "write an application on halfday leave",
            "write a poem",
            "write funny poem",
            "write a resume",
            "write poetry"
        ]

        if any(keyword in user_input for keyword in content_keywords):
            Content(user_input)
        else:
            print(ChatBot(user_input))
